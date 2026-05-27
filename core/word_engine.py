from docx import Document
from docx.shared import Inches
import os
import tempfile
from utils.file_utils import get_output_path


class WordEngine:
    """Generate Word documents from extracted PDF content."""

    @staticmethod
    def create_from_pages(pages, output_path=None, input_path=None, ocr_texts=None):
        """Create a Word document from extracted page data.

        pages: list of dicts with 'text', 'images', 'page_num' from PdfEngine.extract_text_and_images
        ocr_texts: optional dict mapping page_num -> OCR text (for scanned pages)
        """
        doc = Document()

        for i, page in enumerate(pages):
            page_num = page['page_num']
            text = page['text']

            if (not text or len(text.strip()) < 20) and ocr_texts and page_num in ocr_texts:
                text = ocr_texts[page_num]

            if text and text.strip():
                paragraphs = text.split('\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())

            for img_data in page.get('images', []):
                try:
                    ext = img_data['ext']
                    img_bytes = img_data['bytes']
                    with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as tmp:
                        tmp.write(img_bytes)
                        tmp_path = tmp.name
                    doc.add_picture(tmp_path, width=Inches(5.5))
                    os.unlink(tmp_path)
                except Exception:
                    pass

            if i < len(pages) - 1:
                doc.add_page_break()

        if not output_path:
            output_path = get_output_path(input_path or 'output', '.docx')
        doc.save(output_path)
        return output_path
